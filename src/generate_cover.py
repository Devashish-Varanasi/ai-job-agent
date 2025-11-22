# src/generate_cover.py
from typing import Dict, Optional
import os
import json
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Import config with defaults for CI/testing
try:
    from config import LOCAL_LLM_MODEL_PATH, COVER_LETTERS_DIR
except ImportError:
    LOCAL_LLM_MODEL_PATH = "models/orca-mini-3b-gguf2-q4_0.gguf"
    COVER_LETTERS_DIR = "cover_letters"

# Try to use GPT4All if available, otherwise use a template fallback
def generate_cover_with_template(resume_summary: Dict, job: Dict) -> str:
    """Generate a formal, personalized cover letter using template - STRICTLY based on exact resume details"""
    title = job.get("title", "Position")
    company = job.get("company", "Company")
    desc = job.get("description", "")[:1000]  # Reduced - focus more on resume
    
    # STRICTLY use exact resume data - no fallback to placeholders
    if isinstance(resume_summary, str):
        raise ValueError("resume_summary must be a dictionary with parsed resume data")
    
    # Extract EXACT resume details - these are mandatory
    skills_list = resume_summary.get("skills", [])
    if not skills_list:
        print("  Warning: No skills found in resume")
        skills_list = []
    
    all_skills = skills_list  # Use exact skills from resume
    # Get top 6-8 most relevant skills for emphasis
    skills_str = ", ".join(skills_list[:8]) if skills_list else "technical skills"
    
    # Get EXACT name, email, phone from resume
    name = resume_summary.get("name")
    if not name or name.strip() == "":
        raise ValueError("Resume name is required. Please ensure your resume contains a name.")
    
    email = resume_summary.get("email") or ""
    phone = resume_summary.get("phone") or ""
    
    # Get experience snippets - STRICTLY from resume content
    exp_snippets = resume_summary.get("experience_snippets", [])
    education_snippets = resume_summary.get("education_snippets", [])
    
    # Extract meaningful experience (skip headers and very short lines)
    meaningful_exp = [e.strip() for e in exp_snippets if len(e.strip()) > 20 and not e.isupper()] if exp_snippets else []
    
    # Extract job requirements but keep minimal - RESUME is the focus
    desc_lower = desc.lower() if desc else ""
    
    # Find matching skills between resume and job description
    job_keywords = []
    for skill in all_skills:
        skill_lower = skill.lower()
        if skill_lower in desc_lower:
            job_keywords.append(skill)
    
    # Build resume-focused, company-personalized content
    date_str = datetime.now().strftime("%B %d, %Y")
    
    # INTRO - Personalized for this specific company, using exact resume skills
    intro = f"Dear Hiring Manager at {company},\n\nI am writing to express my strong interest in the {title} position at {company}. My expertise in {skills_str} positions me well to contribute effectively to your team's objectives."
    
    # BODY 1 - EXACT RESUME EXPERIENCE - Personalized for company
    if meaningful_exp and len(meaningful_exp) >= 2:
        # Use actual resume content with natural flow
        exp_line1 = meaningful_exp[0]
        exp_line2 = meaningful_exp[1] if len(meaningful_exp) > 1 else ""
        body1 = f"\n\n{exp_line1}"
        if exp_line2:
            body1 += f" Additionally, {exp_line2.lower()}"
        body1 += f" This experience, combined with my proficiency in {skills_str}, has equipped me with the technical expertise and problem-solving capabilities that would be valuable in this role at {company}."
    elif meaningful_exp and len(meaningful_exp) == 1:
        body1 = f"\n\n{meaningful_exp[0]} This hands-on experience with {skills_str} has prepared me to tackle the challenges and opportunities presented by the {title} position at {company}."
    else:
        # Fallback if no experience snippets - still use exact skills
        body1 = f"\n\nMy professional background includes substantial experience working with {skills_str}. I have consistently demonstrated strong technical capabilities and analytical skills, which I am eager to apply to the initiatives and projects at {company}."
    
    # BODY 2 - MORE EXACT SKILLS FROM RESUME - Personalized for company
    skill_highlights = ", ".join(all_skills[:10]) if len(all_skills) > 8 else skills_str
    
    # Add education if available
    education_text = ""
    if education_snippets and len(education_snippets) > 0:
        edu_lines = [e.strip() for e in education_snippets[:2] if len(e.strip()) > 15]
        if edu_lines:
            education_text = f" My educational background includes {edu_lines[0].lower()}. "
    
    body2 = f"\n\nMy core competencies include {skill_highlights}.{education_text}I have developed and refined these skills through practical application and continuous learning."
    
    # Add minimal job connection if skills match
    if job_keywords and len(job_keywords) > 0:
        matched_skills = ", ".join(job_keywords[:4])
        body2 += f" I noticed that the {title} role at {company} requires proficiency in areas such as {matched_skills}, which directly aligns with my technical background and experience. "
    else:
        body2 += f" "
    
    body2 += f"I am confident that my technical expertise, combined with my commitment to excellence, will enable me to make meaningful contributions to {company} and support your team's success."
    
    # CLOSING - Personalized for this specific company
    closing = f"\n\nI am enthusiastic about the opportunity to bring my skills and experience to {company}. I would welcome the chance to discuss how my background aligns with your team's needs and how I can contribute to your organization's continued growth. Thank you for considering my application, and I look forward to the possibility of speaking with you soon.\n\nSincerely,\n{name}"
    
    if email:
        closing += f"\n{email}"
    if phone:
        closing += f"\n{phone}"
    
    return date_str + "\n\n" + intro + body1 + body2 + closing

def generate_cover_gpt4all(resume_summary: Dict, job: Dict, n_tokens: int = 400) -> str:
    # lazy import to avoid making it required
    try:
        from gpt4all import GPT4All
    except Exception as e:
        print("gpt4all not available:", e)
        return generate_cover_with_template(resume_summary, job)

    model_path = LOCAL_LLM_MODEL_PATH
    if not os.path.exists(model_path):
        print("Local LLM model file not found at", model_path)
        return generate_cover_with_template(resume_summary, job)

    try:
        prompt = build_prompt(resume_summary, job)
        gptj = GPT4All(model_path)
        # Using .generate; API depends on model type/version. This is a simple example.
        reply = gptj.generate(prompt, max_tokens=n_tokens)
        if isinstance(reply, (list, tuple)):
            reply = reply[0]
        # Ensure we return a string
        return str(reply) if reply else generate_cover_with_template(resume_summary, job)
    except Exception as e:
        print("LLM generation failed:", e)
        return generate_cover_with_template(resume_summary, job)

def build_prompt(resume_summary: Dict, job: Dict) -> str:
    # STRICTLY use exact resume data - no placeholders
    if isinstance(resume_summary, str):
        raise ValueError("resume_summary must be a dictionary with parsed resume data")
    
    # Extract EXACT resume details - these are mandatory
    name = resume_summary.get("name")
    if not name or name.strip() == "":
        raise ValueError("Resume name is required. Please ensure your resume contains a name.")
    
    skills_list = resume_summary.get("skills", [])
    skills = ", ".join(skills_list[:12]) if skills_list else "technical skills"  # More skills from resume
    
    exp_snippets = resume_summary.get("experience_snippets", [])
    edu_snippets = resume_summary.get("education_snippets", [])
    
    # Use exact experience from resume
    experience = "\n".join([e.strip() for e in exp_snippets[:10] if len(e.strip()) > 15 and not e.isupper()]) if exp_snippets else "Professional experience in relevant field"
    
    # Use exact education from resume
    education = "\n".join([e.strip() for e in edu_snippets[:5] if len(e.strip()) > 10]) if edu_snippets else ""
    
    email = resume_summary.get("email") or ""
    phone = resume_summary.get("phone") or ""
    
    job_title = job.get("title", "")
    company = job.get("company", "")
    job_desc = job.get("description", "")[:600]  # Reduced job description - focus on resume
    
    contact_info = ""
    if email:
        contact_info += f"\nEmail: {email}"
    if phone:
        contact_info += f"\nPhone: {phone}"
    
    prompt = (
        f"You are a professional career advisor writing a personalized cover letter for {company} that STRICTLY uses ONLY the candidate's resume information.\n\n"
        f"CANDIDATE'S RESUME - USE EXACT INFORMATION ONLY (DO NOT INVENT ANYTHING):\n"
        f"Name: {name}{contact_info}\n"
        f"Skills (use exact skills listed): {skills}\n\n"
        f"Professional Experience (use exact experience details):\n{experience}\n\n"
        f"Education (use if available):\n{education}\n\n"
        f"JOB POSTING (for context only - mention company name naturally):\n"
        f"Position: {job_title}\n"
        f"Company: {company}\n"
        f"Brief Description: {job_desc[:400]}\n\n"
        f"INSTRUCTIONS FOR WRITING THE COVER LETTER:\n"
        f"1. DATE: Start with today's date ({datetime.now().strftime('%B %d, %Y')})\n"
        f"2. GREETING: 'Dear Hiring Manager at {company},' (Always include the specific company name)\n"
        f"3. OPENING (2-3 sentences): Express genuine interest in the {job_title} position at {company}. " 
        f"Immediately mention 3-4 EXACT SKILLS from the candidate's resume (use the exact skill names listed above). " 
        f"Personalize the opening for {company} specifically.\n"
        f"4. BODY PARAGRAPH 1 (4-5 sentences): Describe the candidate's ACTUAL EXPERIENCE from their resume above. " 
        f"Use the EXACT experience details provided. Paraphrase the experience naturally but use only the facts from their resume. " 
        f"Connect how this experience would benefit {company} specifically. Mention {company} by name at least once in this paragraph.\n"
        f"5. BODY PARAGRAPH 2 (3-4 sentences): Showcase MORE EXACT SKILLS from the candidate's resume. " 
        f"List specific technical competencies mentioned in their skills section. " 
        f"If education is provided, mention it naturally. " 
        f"Explain how these qualifications align with the needs at {company}. Personalize for {company}.\n"
        f"6. CLOSING (2-3 sentences): Express enthusiasm for contributing to {company} specifically. " 
        f"Mention the company name naturally. Thank them and express availability for an interview.\n"
        f"7. SIGNATURE: 'Sincerely,' followed by the candidate's exact name: {name}\n"
        f"8. CONTACT INFO: If email and phone are provided, include them after the signature\n\n"
        f"CRITICAL REQUIREMENTS - FOLLOW STRICTLY:\n"
        f"- Use ONLY information from the candidate's resume above - DO NOT invent anything\n"
        f"- Use the candidate's EXACT name: {name} (never use placeholders)\n"
        f"- Use the EXACT skills listed in the skills section above\n"
        f"- Use the EXACT experience details provided in the experience section\n"
        f"- Personalize each cover letter specifically for {company} - mention the company name 3-4 times naturally\n"
        f"- The cover letter should be 85-95% about the CANDIDATE'S resume content\n"
        f"- Only 5-15% should reference the job description for context\n"
        f"- Do NOT make up any experience, skills, achievements, or education not in the resume\n"
        f"- Professional, confident tone grounded in actual resume facts\n"
        f"- Length: 350-450 words\n"
        f"- Each cover letter should be unique and personalized for {company}, but based entirely on the resume\n"
    )
    return prompt

def generate_cover(resume_summary: Dict, job: Dict, prefer_local_llm: bool = True) -> str:
    try:
        if prefer_local_llm:
            return generate_cover_gpt4all(resume_summary, job)
        else:
            return generate_cover_with_template(resume_summary, job)
    except Exception as e:
        print(f"Error generating cover letter: {e}")
        # Fallback to template
        try:
            return generate_cover_with_template(resume_summary, job)
        except Exception as e2:
            print(f"Error with fallback template: {e2}")
            return "Failed to generate cover letter due to technical issues."

def save_cover_letter_docx(cover_text: str, job: Dict, resume_summary: Dict, output_dir: Optional[str] = None) -> str:
    """Save cover letter as a formatted DOCX file"""
    if output_dir is None:
        output_dir = COVER_LETTERS_DIR
    
    # Create directory if it doesn't exist
    os.makedirs(output_dir, exist_ok=True)
    
    # Generate filename
    company = job.get("company", "Company").replace("/", "-").replace("\\", "-")
    title = job.get("title", "Position").replace("/", "-").replace("\\", "-")
    job_id = job.get("id", "unknown")
    filename = f"CoverLetter_{company}_{title}_{job_id}.docx"
    filepath = os.path.join(output_dir, filename)
    
    # Create document
    doc = Document()
    
    # Set up default paragraph formatting
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
    
    # Get candidate info - STRICTLY from resume (no placeholders)
    if isinstance(resume_summary, str):
        raise ValueError("resume_summary must be a dictionary with parsed resume data")
    
    # Extract EXACT resume details
    name = resume_summary.get("name")
    if not name or name.strip() == "":
        raise ValueError("Resume name is required for cover letter header")
    
    email = resume_summary.get("email") or ""
    phone = resume_summary.get("phone") or ""
    
    # Add header with contact info - each on separate line, right-aligned
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header_para.add_run(name)
    header_run.font.size = Pt(11)
    
    if email and email.strip():
        email_para = doc.add_paragraph()
        email_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        email_run = email_para.add_run(email)
        email_run.font.size = Pt(11)
    
    if phone and phone.strip():
        phone_para = doc.add_paragraph()
        phone_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        phone_run = phone_para.add_run(phone)
        phone_run.font.size = Pt(11)
    
    # Add date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_run = date_para.add_run(datetime.now().strftime("%B %d, %Y"))
    date_run.font.size = Pt(11)
    
    # Add space
    doc.add_paragraph()
    
    # Add recipient info
    recipient_para = doc.add_paragraph()
    recipient_run = recipient_para.add_run(f"Hiring Manager\n{job.get('company', 'Company')}")
    recipient_run.font.size = Pt(11)
    
    # Add space
    doc.add_paragraph()
    
    # Parse and add cover letter body
    paragraphs = cover_text.split('\n\n')
    for para_text in paragraphs:
        if para_text.strip():
            # Skip date if already in header
            if datetime.now().strftime("%B %d, %Y") in para_text:
                continue
            
            # Check if this paragraph contains the signature section
            para_text_lower = para_text.lower()
            if "sincerely" in para_text_lower:
                # Handle signature section - split by single newlines to keep name together
                lines = para_text.strip().split('\n')
                in_closing_text = True
                
                for i, line in enumerate(lines):
                    line = line.strip()
                    if not line:
                        continue
                    
                    line_lower = line.lower()
                    
                    # Check if this is the "Sincerely," line
                    if line_lower.startswith("sincerely"):
                        # Add extra space before signature
                        doc.add_paragraph()  # Blank line for spacing
                        sig_para = doc.add_paragraph(line)
                        sig_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        para_format = sig_para.paragraph_format
                        para_format.space_after = Pt(6)
                        in_closing_text = False  # Now we're in signature area
                    elif not in_closing_text:
                        # This is name, email, or phone - add as separate left-aligned paragraph
                        # Keep each line on its own paragraph to prevent wrapping
                        sig_line = doc.add_paragraph(line)
                        sig_line.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        para_format = sig_line.paragraph_format
                        para_format.space_after = Pt(3)  # Small space between signature lines
                        # Set font size for signature lines
                        for run in sig_line.runs:
                            run.font.size = Pt(11)
                    else:
                        # This is still closing text before "Sincerely"
                        para = doc.add_paragraph(line)
                        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        para_format = para.paragraph_format
                        para_format.space_after = Pt(6)
                        para_format.line_spacing = 1.15
            else:
                # Regular body paragraph
                para = doc.add_paragraph(para_text.strip())
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para_format = para.paragraph_format
                para_format.space_after = Pt(6)
                para_format.line_spacing = 1.15
    
    # Save document
    doc.save(filepath)
    print(f"  Saved cover letter: {filename}")
    return filepath